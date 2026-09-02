#!/usr/bin/env python3
"""Генерация инструкции для импорта в Google Документы."""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# Палитра
ACCENT = RGBColor(0x1A, 0x73, 0xE8)       # синий Google
ACCENT_DARK = RGBColor(0x17, 0x4E, 0xA6)
TEXT = RGBColor(0x20, 0x21, 0x24)
MUTED = RGBColor(0x5F, 0x63, 0x68)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RPE_COLORS = {
    10: "FCE8E6",
    9: "FEF0E6",
    8: "FFF8E1",
    7: "E8F5E9",
}


def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc, color="DADCE0"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_paragraph(paragraph, size=11, bold=False, color=TEXT, space_after=8, space_before=0):
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.line_spacing = 1.25
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color


def add_section_heading(doc, number: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)

    num = p.add_run(f"{number}. ")
    num.font.name = "Arial"
    num.font.size = Pt(14)
    num.font.bold = True
    num.font.color.rgb = ACCENT

    t = p.add_run(title)
    t.font.name = "Arial"
    t.font.size = Pt(14)
    t.font.bold = True
    t.font.color.rgb = TEXT


def add_body(doc, text: str, indent_cm=0):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    return p


def add_sub_item(doc, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3

    lb = p.add_run(f"{label} ")
    lb.font.name = "Arial"
    lb.font.size = Pt(11)
    lb.font.bold = True
    lb.font.color.rgb = ACCENT_DARK

    body = p.add_run(text)
    body.font.name = "Arial"
    body.font.size = Pt(11)
    body.font.color.rgb = TEXT


def add_rpe_table(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(2.2), Cm(13.5)]
    headers = [("RPE", "Описание")]
    rows = [
        ("10", "нет запаса вообще"),
        ("9", "1 повтор или 2-3 сек в запасе."),
        ("8", "2 повтора или 4-6 сек в запасе."),
        ("7", "3 возможно 4 или 7-10сек и т.д."),
    ]

    # заголовок
    for i, (h1, h2) in enumerate([headers[0]]):
        row = table.rows[0]
        for j, val in enumerate([h1, h2]):
            cell = row.cells[j]
            cell.width = widths[j]
            set_cell_shading(cell, "E8F0FE")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = ACCENT_DARK

    for idx, (rpe, desc) in enumerate(rows, start=1):
        row = table.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width, c1.width = widths[0], widths[1]
        set_cell_shading(c0, RPE_COLORS.get(int(rpe), "FFFFFF"))
        set_cell_shading(c1, "FAFAFA")

        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(rpe)
        r0.font.name = "Arial"
        r0.font.size = Pt(12)
        r0.font.bold = True
        r0.font.color.rgb = TEXT

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        r1.font.color.rgb = TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build_document() -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("Инструкции")
    tr.font.name = "Arial"
    tr.font.size = Pt(28)
    tr.font.bold = True
    tr.font.color.rgb = TEXT

    add_horizontal_line(doc)

    # 1. Разминка
    add_section_heading(doc, "1", "Разминка")
    add_sub_item(
        doc,
        "а.",
        'Перед лазанием желательно использовать полный блок с листа "разминка" но он имеет '
        "рекомендательный характер, упражнения можно заменять. НО висы и разминку пальцев делать!",
    )
    add_sub_item(
        doc,
        "б.",
        "Пред силовыми разминку производим локально, весь блок делать не надо. Т.е. Если делаем "
        "упражнения на верх, то делаем немного общей разминки по верху тела, а потом упражнение которое "
        "собираемся делать в облегченном формате. Например: планируем подтягивания 30кг*5. разминка: "
        "общая на верх тела(чуть мобильности, работа по плечам и ротаторам плеча 1-2 подхода, без утомления) "
        "затем подтягивания без веса 5-8 раз, с 10-15 кг 5 раз, с 20-25 кг 3р и только потом 30кг*5 раз. "
        "По той же схеме пальцы и ноги.",
    )

    add_horizontal_line(doc, "E8EAED")

    # 2. RPE
    add_section_heading(doc, "2", "RPE-")
    add_body(
        doc,
        "Это субъективная оценка количества повторов/секунд в запасе после выполнения подхода.",
    )
    add_rpe_table(doc)
    add_body(
        doc,
        "Пример. Подтянулся 5 раз с 20 кг,все ровно и красиво, в конце оцениваешь сколько еще повторов "
        "смог бы сделать. И, предположим, 2 повтора еще можно выполнить, значит RPE 8.",
        indent_cm=0.3,
    )

    add_horizontal_line(doc, "E8EAED")

    # 4. Обозначения
    add_section_heading(doc, "4", "Обозначения A1,A2. B1,B2.")
    add_body(
        doc,
        "Это значит, что упражнения с одинаковой буквой выполняются по очереди, а цифра обозначает какое "
        "упражнение первое, а какое 2-е/3-е. Упражнения обозначенные одинаковой буквой и цифрой( A1,A1) "
        "выполняются так: выполнить все подходы упражнения, что расположено выше в таблице, затем то, что ниже.",
    )

    add_horizontal_line(doc, "E8EAED")

    # 5. Отдых
    add_section_heading(doc, "5", "Отдых в упражнениях с обозначением A1,A2 и т.д.")
    add_body(
        doc,
        "Интервал отдыха A1,A2 2мин. Это означает что суммарное время отдыха на 2 упражнения 2 минуты. "
        "После первого упражнения отдых 1 минута. Затем, второе упражнение и еще минута отдыха.",
    )

    return doc


def main():
    doc = build_document()
    out = "/Users/maksimizrailev/Documents/hexlet_studying/ski-guide-Italy/Инструкции_тренировки.docx"
    doc.save(out)

    desktop = "/Users/maksimizrailev/Desktop/Инструкции_тренировки.docx"
    import shutil
    shutil.copy(out, desktop)
    print(out)
    print(desktop)


if __name__ == "__main__":
    main()
